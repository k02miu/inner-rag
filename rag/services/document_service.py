import io
import logging
import os
import re
import tempfile

import requests
from bs4 import BeautifulSoup
from docx import Document
from PyPDF2 import PdfReader

logger = logging.getLogger(__name__)


class DocumentService:
    """
    ドキュメントサービス
    """

    def __init__(self) -> None:
        pass

    def extract_text(self, file_content: bytes, file_type: str) -> str | None:
        """
        ファイルからテキストを抽出する
        ファイルの種類（pdf, docx, xlsx等）
        """
        try:
            if file_type in ["pdf"]:
                return self._extract_from_pdf(file_content)
            elif file_type in ["docx", "doc"]:
                return self._extract_from_docx(file_content)
            elif file_type in ["xlsx", "xls"]:
                return self._extract_from_excel(file_content)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return None
        except Exception as e:
            logger.error(
                f"Error extracting text from {file_type} file: {str(e)}"
            )
            return None

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """
        PDFファイルからテキストを抽出する
        """
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)

        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return "\n\n".join(text_parts)

    def _extract_from_docx(self, file_content: bytes) -> str:
        """
        DOCXファイルからテキストを抽出する
        """
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)

        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_parts.append(paragraph.text)

        # テーブルからもテキストを抽出
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    if cell.text:
                        row_texts.append(cell.text)
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        return "\n\n".join(text_parts)

    def _extract_from_excel(self, file_content: bytes) -> str:
        """
        Excelファイルからテキストを抽出する
        """
        # pandas を使用してExcelファイルを読み込む
        # ここでは一時ファイルに保存してから読み込む
        try:
            import pandas as pd

            with tempfile.NamedTemporaryFile(
                delete=False, suffix=".xlsx"
            ) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name

            # Excelファイルを読み込む
            text_parts = []
            excel_file = pd.ExcelFile(temp_file_path)

            # 各シートを処理
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name)
                text_parts.append(f"Sheet: {sheet_name}")
                text_parts.append(df.to_string(index=False))

            # 一時ファイルを削除
            os.unlink(temp_file_path)

            return "\n\n".join(text_parts)
        except ImportError:
            logger.warning(
                "pandas is not installed. Excel extraction will be limited."
            )
            return (
                "Excel file content could not be extracted."
                " Please install pandas for Excel support."
            )
        except Exception as e:
            logger.error(f"Error extracting text from Excel file: {str(e)}")
            return f"Error extracting Excel content: {str(e)}"

    def extract_from_url(self, url: str) -> str | None:
        """
        URLからコンテンツを抽出する
        """
        try:
            # URLからコンテンツを取得
            headers = {
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
                    " AppleWebKit/537.36 "
                    "(KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
                )
            }
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()

            # コンテンツタイプを確認
            content_type = response.headers.get("Content-Type", "").lower()

            # HTMLの場合
            if "text/html" in content_type:
                return self._extract_from_html(response.text, url)
            # PDFの場合
            elif "application/pdf" in content_type:
                return self._extract_from_pdf(response.content)
            # DOCXの場合
            elif (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ) in content_type:
                return self._extract_from_docx(response.content)
            # その他のテキストベースのコンテンツ
            elif "text/" in content_type:
                return response.text
            else:
                logger.warning(f"Unsupported content type: {content_type}")
                return None
        except requests.RequestException as e:
            logger.error(f"Error fetching URL {url}: {str(e)}")
            return None
        except Exception as e:
            logger.error(f"Error extracting content from URL {url}: {str(e)}")
            return None

    def _extract_from_html(self, html_content: str, url: str) -> str:
        """
        HTMLからテキストを抽出する
        """
        try:
            soup = BeautifulSoup(html_content, "html.parser")

            # 不要な要素を削除
            for element in soup(
                ["script", "style", "nav", "footer", "header"]
            ):
                element.decompose()

            # タイトルを取得
            title = soup.title.string if soup.title else "No Title"

            # メインコンテンツを抽出
            main_content = ""

            # article, main, または div.content
            # などの主要コンテンツを含む要素を探す
            main_elements = soup.select(
                "article, main, .content, #content, .main, #main"
            )
            if main_elements:
                for element in main_elements:
                    main_content += (
                        element.get_text(separator="\n", strip=True) + "\n\n"
                    )
            else:
                # 主要コンテンツが見つからない場合は、bodyから抽出
                main_content = (
                    soup.body.get_text(separator="\n", strip=True)
                    if soup.body
                    else ""
                )

            # テキストを整形
            # 複数の空白を1つに置換
            main_content = re.sub(r"\s+", " ", main_content)
            # 複数の改行を1つに置換
            main_content = re.sub(r"\n+", "\n", main_content)

            return f"Title: {title}\n\nURL: {url}\n\n{main_content}"
        except Exception as e:
            logger.error(f"Error extracting text from HTML: {str(e)}")
            return f"Error extracting HTML content: {str(e)}"

    def chunk_text(
        self, text: str, chunk_size: int = 1000, overlap: int = 100
    ) -> list[str]:
        """
        テキストをチャンク（断片）に分割する
        """
        # テキストが短い場合はそのまま返す
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            # チャンクのサイズを決定
            end = start + chunk_size

            # テキストの末尾に達した場合
            if end >= len(text):
                chunks.append(text[start:])
                break

            # 段落や文の境界で区切るように調整
            if text[end] != " " and end < len(text) - 1:
                # 空白を探して、そこで区切る
                next_space = text.find(" ", end)
                next_newline = text.find("\n", end)

                if next_newline != -1 and (
                    next_space == -1 or next_newline < next_space
                ):
                    end = next_newline + 1
                elif next_space != -1:
                    end = next_space + 1

            # チャンクを追加
            chunks.append(text[start:end])

            # 次のチャンクの開始位置を設定（オーバーラップを考慮）
            start = end - overlap

        return chunks
